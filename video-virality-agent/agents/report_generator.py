"""
DOCX report generator — content-team focused, black text, all key sections.
"""

import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# All body text is black. Only headings and labels use colour.
C_BLACK  = RGBColor(0, 0, 0)
C_PURPLE = RGBColor(80, 60, 200)
C_RED    = RGBColor(180, 0, 0)
C_GREEN  = RGBColor(0, 130, 60)
C_YELLOW = RGBColor(160, 120, 0)
C_CYAN   = RGBColor(0, 100, 160)
C_MUTED  = RGBColor(80, 80, 80)
C_WHITE  = RGBColor(255, 255, 255)


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _grade_color(pct: int) -> RGBColor:
    if pct >= 70: return C_GREEN
    if pct >= 50: return C_YELLOW
    if pct >= 35: return RGBColor(180, 100, 0)
    return C_RED


def _h(doc, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = C_PURPLE
        run.font.size = Pt(14)
        run.bold = True
    elif level == 2:
        run.font.color.rgb = C_CYAN
        run.font.size = Pt(11)
        run.bold = True
    else:
        run.font.color.rgb = C_MUTED
        run.font.size = Pt(10)
        run.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def _bullet(doc, text: str, color: RGBColor = None, prefix: str = "•"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"{prefix}  {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = color if color else C_BLACK


def _row(doc, label: str, value: str, vc: RGBColor = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    lbl = p.add_run(f"{label}: ")
    lbl.bold = True
    lbl.font.size = Pt(10)
    lbl.font.color.rgb = C_MUTED
    val = p.add_run(str(value))
    val.font.size = Pt(10)
    val.font.color.rgb = vc if vc else C_BLACK


def _add_page_watermark_footer(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run("Video Virality Analysis by Social Media Team   |   " +
                         datetime.now().strftime("%d %b %Y"))
        run.font.size = Pt(8)
        run.font.color.rgb = C_MUTED
        run.font.italic = True


def _score_table(doc, breakdown: dict):
    labels = {
        "hook": "Hook", "story": "Story", "script": "Script",
        "audio": "Audio", "visual": "Visual", "editing": "Editing",
        "retention": "Retention", "thumbnail": "Thumbnail",
        "title": "Title", "virality": "Virality",
    }
    rows = [(labels.get(k, k.replace("_", " ").title()), v)
            for k, v in breakdown.items()]
    if not rows:
        return

    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for cell, txt in zip(table.rows[0].cells, ["Dimension", "Score", "Rating"]):
        cell.text = txt
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        _set_cell_bg(cell, "503CC8")
        cell.paragraphs[0].runs[0].font.color.rgb = C_WHITE

    for i, (label, score) in enumerate(rows):
        row = table.rows[i + 1].cells
        if score is None:
            row[0].text = label
            row[0].paragraphs[0].runs[0].font.size = Pt(9)
            row[1].text = "N/A"
            row[1].paragraphs[0].runs[0].font.size = Pt(9)
            row[1].paragraphs[0].runs[0].font.color.rgb = C_MUTED
            row[2].text = "Not uploaded"
            row[2].paragraphs[0].runs[0].font.size = Pt(9)
            row[2].paragraphs[0].runs[0].font.color.rgb = C_MUTED
            continue
        score_int = int(score)
        is_10 = score_int <= 10
        display = f"{score_int}/10" if is_10 else f"{score_int}/100"
        pct = score_int * 10 if is_10 else score_int
        gc = _grade_color(pct)
        verdict = "Excellent" if pct >= 70 else "Average" if pct >= 50 else "Weak" if pct >= 35 else "Poor"
        row[0].text = label
        row[0].paragraphs[0].runs[0].font.size = Pt(9)
        row[1].text = display
        row[1].paragraphs[0].runs[0].font.size = Pt(9)
        row[1].paragraphs[0].runs[0].bold = True
        row[1].paragraphs[0].runs[0].font.color.rgb = gc
        row[2].text = verdict
        row[2].paragraphs[0].runs[0].font.size = Pt(9)
        row[2].paragraphs[0].runs[0].font.color.rgb = gc

    doc.add_paragraph()


def generate_report(data: dict, mode: str = "url") -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Normal"].font.color.rgb = C_BLACK

    meta  = data.get("metadata", {})
    ca    = data.get("content_analysis") or data.get("video_analysis") or {}
    ta    = data.get("thumbnail_analysis", {})
    comp  = data.get("competitor_analysis", {})
    comps = data.get("competitors", [])

    if not ca and data.get("virality_score") is not None:
        ca = data

    score = int(ca.get("virality_score") or ta.get("score") or 0)
    grade = ca.get("grade") or ta.get("grade") or "F"
    title = meta.get("title") or "Uploaded Video"
    gc    = _grade_color(score)

    # ── COVER ──────────────────────────────────────────────────────────────────
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover.add_run("VIDEO VIRALITY ANALYSIS REPORT")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = C_PURPLE

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = p.add_run(title)
    tr.font.size = Pt(13)
    tr.font.color.rgb = C_BLACK
    tr.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"Score: {score}/100   |   Grade {grade}")
    r2.bold = True
    r2.font.size = Pt(14)
    r2.font.color.rgb = gc

    if meta:
        doc.add_paragraph()
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ch = p3.add_run(f"Channel: {meta.get('channel','—')}   |   "
                        f"Views: {meta.get('view_count', 0):,}   |   "
                        f"Likes: {meta.get('like_count', 0):,}")
        ch.font.size = Pt(10)
        ch.font.color.rgb = C_MUTED

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = p4.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y  %H:%M')}")
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = C_MUTED

    doc.add_page_break()

    # ── 1. SCORE BREAKDOWN ─────────────────────────────────────────────────────
    _h(doc, "1. Score Breakdown")
    if ca.get("executive_summary"):
        p = doc.add_paragraph(ca["executive_summary"])
        p.paragraph_format.space_after = Pt(8)
        if p.runs:
            p.runs[0].font.color.rgb = C_BLACK

    if ca.get("confidence_level"):
        cl = ca["confidence_level"]
        clr = C_GREEN if cl == "High" else C_YELLOW if cl == "Medium" else C_RED
        _row(doc, "Analysis Confidence", cl, clr)

    if ca.get("final_verdict"):
        p = doc.add_paragraph()
        fv = p.add_run(f"Verdict: {ca['final_verdict']}")
        fv.font.color.rgb = C_RED
        fv.bold = True
        fv.font.size = Pt(10)

    doc.add_paragraph()
    if ca.get("breakdown"):
        _score_table(doc, ca["breakdown"])

    # ── 2. PERFORMANCE PREDICTION ─────────────────────────────────────────────
    pp = ca.get("performance_prediction", {})
    if pp:
        _h(doc, "2. Performance Prediction After Upload")
        note = doc.add_paragraph(
            "These are honest estimates based on your content quality score. "
            "Actual numbers depend on your channel size and promotion."
        )
        if note.runs:
            note.runs[0].font.color.rgb = C_MUTED
            note.runs[0].italic = True
        note.paragraph_format.space_after = Pt(8)

        table = doc.add_table(rows=6, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        pairs = [
            ("Views in 30 Days",          pp.get("predicted_views_30_days", "—")),
            ("Avg View Duration (AVD)",    pp.get("predicted_avg_view_duration", "—")),
            ("Watch Time (30 Days)",       pp.get("predicted_watch_time_hours", "—")),
            ("New Subscribers",            pp.get("predicted_new_subscribers", "—")),
            ("CTR Estimate",               pp.get("ctr_estimate", "—")),
            ("Confidence Level",           pp.get("confidence", "Medium")),
        ]
        for row_idx, (label, value) in enumerate(pairs):
            row = table.rows[row_idx].cells
            row[0].text = label
            row[0].paragraphs[0].runs[0].bold = True
            row[0].paragraphs[0].runs[0].font.size = Pt(10)
            row[0].paragraphs[0].runs[0].font.color.rgb = C_BLACK
            row[1].text = str(value)
            row[1].paragraphs[0].runs[0].font.size = Pt(10)
            row[1].paragraphs[0].runs[0].font.color.rgb = C_BLACK

        doc.add_paragraph()
        if pp.get("what_will_hurt_most"):
            _bullet(doc, f"Biggest risk: {pp['what_will_hurt_most']}", C_RED, "⚠")
        if pp.get("what_could_boost_it"):
            _bullet(doc, f"Biggest opportunity: {pp['what_could_boost_it']}", C_GREEN, "✅")

    # ── 3. CRITICAL ISSUES ────────────────────────────────────────────────────
    issues = ca.get("critical_issues", [])
    if issues:
        _h(doc, "3. Critical Issues  —  fix these before uploading")
        for ci in issues:
            if isinstance(ci, dict):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.1)
                p.paragraph_format.space_before = Pt(6)
                ts = f" @ {ci['timestamp']}" if ci.get("timestamp") else ""
                lbl = p.add_run(f"⚠  {ci.get('issue', '')}{ts}")
                lbl.bold = True
                lbl.font.color.rgb = C_RED
                lbl.font.size = Pt(10)
                if ci.get("why_it_hurts"):
                    _bullet(doc, f"Why: {ci['why_it_hurts']}", C_BLACK, "  →")
                if ci.get("fix"):
                    _bullet(doc, f"Fix: {ci['fix']}", C_GREEN, "  ✅")
                if ci.get("expected_impact"):
                    _bullet(doc, f"Impact: {ci['expected_impact']}", C_CYAN, "  📈")
            else:
                _bullet(doc, str(ci), C_RED, "⚠")

    # ── 4. QUICK WINS ─────────────────────────────────────────────────────────
    if ca.get("quick_wins"):
        _h(doc, "4. Quick Wins  —  do these first (under 10 minutes each)")
        for w in ca["quick_wins"]:
            _bullet(doc, w, C_GREEN, "⚡")

    # ── 5. TIMESTAMPED OBSERVATIONS ──────────────────────────────────────────
    obs = ca.get("timestamped_observations", [])
    if obs:
        _h(doc, "5. Timestamped Observations")
        for o in obs:
            if isinstance(o, dict):
                sev = o.get("severity", "ok")
                clr = C_RED if sev == "critical" else C_YELLOW if sev == "warning" else C_GREEN
                ico = "🔴" if sev == "critical" else "🟡" if sev == "warning" else "🟢"
                _bullet(doc, f"[{o.get('time', '?')}]  {o.get('observation', '')}", clr, ico)
            else:
                _bullet(doc, str(o), C_BLACK, "•")

    # ── 6. AUDIO ANALYSIS ────────────────────────────────────────────────────
    aa = ca.get("audio_analysis", {})
    if aa:
        _h(doc, "6. Audio Analysis")
        for key, label in [
            ("verdict",             "Overall Verdict"),
            ("filler_word_problem", "Filler Words"),
            ("pacing_verdict",      "Pacing"),
            ("audio_visual_sync",   "A/V Sync"),
            ("script_quality",      "Script Quality"),
        ]:
            if aa.get(key):
                _row(doc, label, aa[key], C_BLACK)

    # ── 7. CONTENT & TITLE ANALYSIS ──────────────────────────────────────────
    _h(doc, "7. Content & Title Analysis")
    if ca.get("strengths"):
        _h(doc, "Strengths", level=2)
        for s in ca["strengths"]:
            _bullet(doc, s, C_GREEN, "✓")
    if ca.get("weaknesses"):
        _h(doc, "Weaknesses", level=2)
        for w in ca["weaknesses"]:
            _bullet(doc, w, C_RED, "✗")
    sa = ca.get("story_arc", {})
    if sa:
        _h(doc, "Story Arc", level=2)
        if sa.get("phase_1_setup"):
            _row(doc, "Setup", sa["phase_1_setup"], C_BLACK)
        if sa.get("phase_2_tension"):
            _row(doc, "Tension", sa["phase_2_tension"], C_BLACK)
        if sa.get("phase_3_resolution"):
            _row(doc, "Resolution", sa["phase_3_resolution"], C_BLACK)
        if sa.get("arc_verdict"):
            vc = C_GREEN if sa["arc_verdict"] == "Strong" else C_RED
            _row(doc, "Arc Verdict", sa["arc_verdict"], vc)
        if sa.get("storytelling_fix"):
            _bullet(doc, f"Fix: {sa['storytelling_fix']}", C_CYAN, "💡")
    if ca.get("ctr_prediction"):
        _row(doc, "CTR Prediction", ca["ctr_prediction"], C_BLACK)
    if ca.get("retention_prediction"):
        _row(doc, "Retention Prediction", ca["retention_prediction"], C_BLACK)
    if ca.get("virality_prediction"):
        _row(doc, "Virality Prediction", ca["virality_prediction"], C_BLACK)

    # ── 8. REWRITTEN HOOK SCRIPT ──────────────────────────────────────────────
    if ca.get("rewritten_hook"):
        _h(doc, "8. Rewritten Hook Script  —  give this word-for-word to the creator")
        p = doc.add_paragraph(ca["rewritten_hook"])
        p.paragraph_format.left_indent = Inches(0.3)
        if p.runs:
            p.runs[0].font.color.rgb = C_BLACK
            p.runs[0].italic = True

    # ── 9. SCRIPT LINE-BY-LINE FIXES ─────────────────────────────────────────
    sll = ca.get("script_line_by_line", [])
    if sll:
        _h(doc, "9. Script Line-by-Line Fixes")
        for entry in sll:
            if not isinstance(entry, dict):
                continue
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.1)
            p.paragraph_format.space_before = Pt(6)
            bad = p.add_run(f'BAD:   "{entry.get("original", "")}"')
            bad.font.color.rgb = C_RED
            bad.font.size = Pt(9)
            bad.bold = True
            if entry.get("problem"):
                prob = doc.add_paragraph()
                prob.paragraph_format.left_indent = Inches(0.25)
                prob_run = prob.add_run(f'Problem: {entry["problem"]}')
                prob_run.font.size = Pt(9)
                prob_run.font.color.rgb = C_BLACK
            good_p = doc.add_paragraph()
            good_p.paragraph_format.left_indent = Inches(0.25)
            good_p.paragraph_format.space_after = Pt(6)
            good = good_p.add_run(f'GOOD:  "{entry.get("rewrite", "")}"')
            good.font.color.rgb = C_GREEN
            good.font.size = Pt(9)
            good.bold = True

    # ── 10. FULL 60-SECOND REWRITTEN SCRIPT ──────────────────────────────────
    full_script = ca.get("full_60s_script", "")
    if full_script:
        _h(doc, "10. Full 60-Second Rewritten Script")
        p = doc.add_paragraph()
        p.add_run("Give this to the scriptwriter and video editor as-is. "
                  "All visual directions are in [brackets].").font.color.rgb = C_MUTED
        p = doc.add_paragraph(full_script)
        p.paragraph_format.left_indent = Inches(0.3)
        if p.runs:
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.color.rgb = C_BLACK

    # ── 11. 10 IMPROVED TITLES ────────────────────────────────────────────────
    titles = ca.get("ten_improved_titles", [])
    if titles:
        _h(doc, "11. 10 Improved Title Options  —  test these for better CTR")
        for i, t in enumerate(titles, 1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(3)
            num = p.add_run(f"  {i:02d}.  ")
            num.bold = True
            num.font.color.rgb = C_PURPLE
            body = p.add_run(str(t))
            body.font.size = Pt(10)
            body.font.color.rgb = C_BLACK

    # ── 12. 5 THUMBNAIL CONCEPTS ──────────────────────────────────────────────
    concepts = ca.get("five_thumbnail_concepts", []) or ta.get("five_thumbnail_concepts", [])
    if concepts:
        _h(doc, "12. 5 Thumbnail Concepts  —  brief your designer with these")
        for i, c in enumerate(concepts, 1):
            p = doc.add_paragraph()
            num = p.add_run(f"Concept {i}:  ")
            num.bold = True
            num.font.color.rgb = C_PURPLE
            num.font.size = Pt(10)
            body = p.add_run(str(c))
            body.font.size = Pt(10)
            body.font.color.rgb = C_BLACK
            p.paragraph_format.space_after = Pt(5)

    # ── 13. FINAL VERDICT ─────────────────────────────────────────────────────
    if ca.get("final_verdict"):
        _h(doc, "13. Final Verdict")
        p = doc.add_paragraph(ca["final_verdict"])
        if p.runs:
            p.runs[0].font.color.rgb = C_RED
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(11)

    # ── 14. UPLOAD STRATEGY & METADATA ───────────────────────────────────────
    us = ca.get("upload_strategy", {})
    if us:
        _h(doc, "14. Upload Strategy & Metadata")
        days = us.get("best_days", [])
        t_utc = us.get("best_time_utc", "")
        if days:
            _row(doc, "Best Days to Post", ", ".join(days), C_GREEN)
        if t_utc:
            _row(doc, "Best Time (UTC)", t_utc, C_BLACK)
        if us.get("reasoning"):
            _bullet(doc, us["reasoning"], C_BLACK, "💡")
        mf = us.get("metadata_fixes", {})
        if mf:
            _h(doc, "Metadata Fixes", level=2)
            if mf.get("title_rewrite"):
                _row(doc, "Rewritten Title", mf["title_rewrite"], C_BLACK)
            if mf.get("description_first_line"):
                _row(doc, "Description Opening", mf["description_first_line"], C_BLACK)
            tags = mf.get("must_add_tags", [])
            if tags:
                _row(doc, "Must-Add Tags", "  ".join(f"#{t}" for t in tags), C_BLACK)
        chapters = us.get("chapter_timestamps", [])
        if chapters:
            _h(doc, "Chapter Timestamps", level=2)
            for ch in chapters:
                _bullet(doc, ch, C_BLACK, "🕐")

    # ── 15. THUMBNAIL ANALYSIS (only when thumbnail was uploaded) ─────────────
    if ta and (ta.get("improvements") or ta.get("weaknesses")):
        _h(doc, "15. Thumbnail Analysis")
        if ta.get("ctr_prediction"):
            _row(doc, "CTR Prediction", ta["ctr_prediction"], C_BLACK)
        for w in ta.get("weaknesses", []):
            _bullet(doc, w, C_RED, "❌")
        for w in ta.get("improvements", []):
            _bullet(doc, w, C_GREEN, "✅")

    # ── 16. COMPETITOR VIDEOS THIS WEEK ──────────────────────────────────────
    if comps:
        _h(doc, "16. Competing Videos This Week")
        for i, c in enumerate(comps, 1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.space_after = Pt(3)
            num = p.add_run(f"{i}. ")
            num.bold = True
            num.font.color.rgb = C_PURPLE
            title_run = p.add_run(f'{c.get("title", "—")} — {c.get("channel", "—")}')
            title_run.font.size = Pt(10)
            title_run.font.color.rgb = C_BLACK
            vc = c.get("view_count", 0)
            if vc:
                vc_run = p.add_run(f"  [{vc:,} views]")
                vc_run.font.size = Pt(9)
                vc_run.font.color.rgb = C_MUTED
            if c.get("youtube_url"):
                p2 = doc.add_paragraph()
                p2.paragraph_format.left_indent = Inches(0.35)
                p2.paragraph_format.space_after = Pt(5)
                link = p2.add_run(f"Watch: {c['youtube_url']}")
                link.font.size = Pt(9)
                link.font.color.rgb = C_CYAN
                link.italic = True

    # ── FOOTER WATERMARK ──────────────────────────────────────────────────────
    _add_page_watermark_footer(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
